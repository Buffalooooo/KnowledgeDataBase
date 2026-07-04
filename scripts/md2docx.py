"""
MD → Word 文档转换工具
将 Markdown 演讲稿转换为排版精美的 Word 文档
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def convert(md_path: str, docx_path: str):
    doc = Document()

    # ── 页面设置 ──
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # ── 读取 MD ──
    with open(md_path, "r", encoding="utf-8") as f:
        text = f.read()

    lines = text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].strip()

        # ── 标题行（# 开头） ──
        if line.startswith("# ") and not line.startswith("## "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line.lstrip("# ").strip())
            run.font.size = Pt(22)
            run.bold = True
            run.font.name = "黑体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
            p.paragraph_format.space_before = Pt(24)
            p.paragraph_format.space_after = Pt(12)

        elif line.startswith("## "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(line.lstrip("## ").strip())
            run.font.size = Pt(18)
            run.bold = True
            run.font.name = "黑体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(8)

        elif line.startswith("### "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(line.lstrip("### ").strip())
            run.font.size = Pt(16)
            run.bold = True
            run.font.name = "黑体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)

        # ── 粗体分隔行 ──
        elif line.startswith("**") and line.endswith("**") and line.count("**") == 2:
            content = line.strip("**").strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(content)
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = "黑体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(6)

        # ── "主讲人" 行 ──
        elif "主讲人" in line and "**" in line:
            content = line.replace("**", "")
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(content)
            run.font.size = Pt(14)
            run.font.name = "等线"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
            p.paragraph_format.space_after = Pt(18)

        # ── 分隔线（---） ──
        elif line == "---":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("━" * 40)
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(180, 180, 180)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)

        # ── 正文段落 ──
        elif line:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Pt(24)
            p.paragraph_format.line_spacing = Pt(28)

            # 解析行内加粗 **文字**
            parts = re.split(r"(\*\*.*?\*\*)", line)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    # 处理引号加粗：开头**粗体**
                    run = p.add_run(part)
                run.font.size = Pt(14)
                run.font.name = "仿宋"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")

        # ── 空行跳过 ──
        else:
            pass  # 空行，跳过

        i += 1

    doc.save(docx_path)
    print(f"✅ Word 文档已生成：{docx_path}")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法：python md2docx.py <输入.md> [输出.docx]")
        sys.exit(1)

    md_file = sys.argv[1]
    if len(sys.argv) >= 3:
        docx_file = sys.argv[2]
    else:
        docx_file = str(Path(md_file).with_suffix(".docx"))

    convert(md_file, docx_file)
